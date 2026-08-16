from crewai.tools import tool
from pathlib import Path
import subprocess

SANDBOX_DIR = Path(__file__).parents[3] / "sandbox"
SANDBOX_DIR.mkdir(parents=True, exist_ok=True)


@tool("List Sandbox Files")
def list_sandbox_files() -> str:
    """
    List the filenames currently in the sandbox directory.

    Returns:
        A newline-separated list of filenames, or a message if the
        sandbox is empty.
    """
    names = sorted(p.name for p in SANDBOX_DIR.iterdir())
    return "\n".join(names) if names else "The sandbox is empty."


@tool("Read Sandbox File")
def read_sandbox_file(filename: str) -> str:
    """
    Read and return the text contents of a file in the sandbox directory.

    Args:
        filename: The name of the file to read (e.g. "solution.py").
    Returns:
        The file's contents, or a message if the file does not exist.
    """
    path = SANDBOX_DIR / filename
    if not path.is_file():
        return f"No such file in the sandbox: {filename}"
    return path.read_text()


@tool("Write Sandbox File")
def write_sandbox_file(filename: str, content: str) -> str:
    """
    Write text to a file in the sandbox directory, replacing any existing
    file with the same name.

    Args:
        filename: The name of the file to write (e.g. "solution.py").
        content: The text content to write.
    Returns:
        A confirmation message.
    """
    path = SANDBOX_DIR / filename
    path.write_text(content)
    return f"Wrote {len(content)} characters to {filename}."


@tool("Run Sandbox Python File")
def run_sandbox_python(filename: str) -> str:
    """
    Execute a Python file from the sandbox directory inside an ephemeral
    Docker container, with the sandbox mounted as the working directory,
    and return whatever the script printed to stdout.

    Args:
        filename: The name of the Python file to run (e.g. "solution.py").
    Returns:
        The text printed to stdout by the executed script.
    """
    result = subprocess.run(
        [
            "docker",
            "run",
            "--rm",
            "-v",
            f"{SANDBOX_DIR}:/workspace",
            "-w",
            "/workspace",
            "coder-ml-sandbox",
            "python",
            filename,
        ],
        capture_output=True,
        text=True,
        timeout=180,
    )
    output_dir = SANDBOX_DIR.parent / "output"
    output_dir.mkdir(parents=True, exist_ok=True)

    # Save raw log file
    log_file = output_dir / "container_execution.log"
    full_output = (result.stdout or "") + (
        "\n" + result.stderr if result.stderr else ""
    )
    log_file.write_text(
        f"--- STDOUT ---\n{result.stdout}\n\n--- STDERR ---\n{result.stderr}"
    )

    # Generate dark-theme terminal output image and save to solution.md
    prompt_cmd = f"PS D:\\coder\\sandbox> python {filename}"
    img_path = output_dir / "terminal_output.png"
    _generate_terminal_image(prompt_cmd, full_output, img_path)

    # Copy any generated plot PNG files from sandbox to output directory
    import shutil

    plot_md_section = ""
    plot_files = [
        p for p in SANDBOX_DIR.glob("*.png") if p.name != "terminal_output.png"
    ]
    if plot_files:
        plot_md_section = "### Generated Visualizations & Plots\n\n"
        for p in plot_files:
            dest = output_dir / p.name
            shutil.copy(p, dest)
            plot_md_section += f"![{p.stem}]({p.name})\n\n"

    solution_md = output_dir / "solution.md"
    solution_md.write_text(
        f"# Code Execution Output\n\n"
        f"### Terminal Output Screenshot\n\n"
        f"![Terminal Output](terminal_output.png)\n\n"
        f"{plot_md_section}"
        f"### Raw Text Log\n\n"
        f"```powershell\n{prompt_cmd}\n{full_output}\n```\n"
    )

    # Generate Word Document solution.docx according to specified formatting
    _generate_word_document(filename, prompt_cmd, full_output, output_dir)

    return result.stdout


def _generate_terminal_image(prompt_cmd: str, raw_output: str, img_path: Path):
    try:
        from PIL import Image, ImageDraw, ImageFont
        import textwrap
    except ImportError:
        return

    # Process and line-wrap output lines so nothing gets cut off
    raw_lines = raw_output.splitlines() if raw_output else ["(No output produced)"]
    formatted_lines = []
    
    # Add prompt line
    formatted_lines.append((prompt_cmd, "prompt"))

    for line in raw_lines:
        # Wrap lines longer than 85 chars
        wrapped = textwrap.wrap(line, width=85) if len(line) > 85 else [line]
        for w in wrapped:
            if "Error" in w or "Traceback" in w or "Exception" in w:
                tag = "error"
            elif w.strip().startswith("[") or w.strip().startswith("{"):
                tag = "data"
            else:
                tag = "text"
            formatted_lines.append((w, tag))

    # 2x Supersampling for razor-sharp rendering in Word
    scale = 2
    font_size = 15 * scale
    line_height = 24 * scale
    padding = 20 * scale
    header_height = 36 * scale

    # Auto-detect best monospace font on Windows / OS
    font_candidates = [
        "C:\\Windows\\Fonts\\consola.ttf",
        "C:\\Windows\\Fonts\\consolab.ttf",
        "C:\\Windows\\Fonts\\lucon.ttf",
        "C:\\Windows\\Fonts\\arial.ttf",
        "consola.ttf",
        "arial.ttf"
    ]
    
    font = None
    for fc in font_candidates:
        try:
            font = ImageFont.truetype(fc, font_size)
            break
        except Exception:
            continue
            
    if font is None:
        font = ImageFont.load_default()

    # Calculate high-res width and height
    max_line_len = max((len(text) for text, _ in formatted_lines), default=45)
    img_width = max(900 * scale, max_line_len * 10 * scale + padding * 2)
    img_height = len(formatted_lines) * line_height + padding * 2

    # Create dark background image (#1e1e1e)
    img = Image.new("RGB", (img_width, img_height), color=(30, 30, 30))
    draw = ImageDraw.Draw(img)

    # Render lines directly from top padding (No header bar, No RGB dots)
    y = padding
    for text, tag in formatted_lines:
        if tag == "prompt":
            color = (229, 192, 123)  # Gold / Yellow
        elif tag == "error":
            color = (244, 71, 71)    # Bright Red
        elif tag == "data":
            color = (97, 175, 239)   # Bright Cyan/Blue
        else:
            color = (220, 220, 220)  # Crisp White/Light Grey
            
        draw.text((padding, y), text, font=font, fill=color)
        y += line_height

    # Downscale smoothly to 1x with LANCZOS for anti-aliased crisp text
    final_img = img.resize((img_width // scale, img_height // scale), Image.Resampling.LANCZOS)
    final_img.save(img_path, dpi=(300, 300))


def _generate_word_document(filename: str, prompt_cmd: str, full_output: str, output_dir: Path):
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return

    doc = Document()

    def format_run(run, font_name="Times New Roman", size_pt=14, bold=False, italic=False, color_rgb=(0,0,0)):
        run.font.name = font_name
        run.font.size = Pt(size_pt)
        run.bold = bold
        run.italic = italic
        run.font.color.rgb = RGBColor(*color_rgb)

    # 1. Document Title (Experiment / <N>) -> Times New Roman, 16 pt, Bold, Centered
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(f"Experiment / Assignment: {filename}")
    format_run(title_run, font_name="Times New Roman", size_pt=16, bold=True)

    doc.add_paragraph()  # Spacing

    # Read code content from sandbox directory
    code_path = SANDBOX_DIR / filename
    code_content = code_path.read_text() if code_path.is_file() else "(Code file not found)"

    def add_section_heading(text: str):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        format_run(run, font_name="Times New Roman", size_pt=16, bold=True)

    def add_body_text(text: str):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        format_run(run, font_name="Times New Roman", size_pt=14, bold=False)

    # Read exact assignment prompt given by user
    assignment_file = output_dir / "assignment.txt"
    if assignment_file.is_file() and assignment_file.read_text().strip():
        aim_text = assignment_file.read_text().strip()
    else:
        aim_text = f"To write and execute Python code for {filename} and record its output and visualizations."

    # 2. Section Headings: Aim, Theory, Software Used, Code, Output -> Times New Roman, 16pt, Bold, Left
    # Body Text: Times New Roman, 14pt, Regular, Left
    add_section_heading("Aim")
    add_body_text(aim_text)

    # Read theory explanation generated by agent
    theory_file_sandbox = SANDBOX_DIR / "theory.txt"
    theory_file_output = output_dir / "theory.txt"
    if theory_file_sandbox.is_file() and theory_file_sandbox.read_text().strip():
        theory_text = theory_file_sandbox.read_text().strip()
    elif theory_file_output.is_file() and theory_file_output.read_text().strip():
        theory_text = theory_file_output.read_text().strip()
    else:
        theory_text = "Implementation of algorithm and computational tasks using standard Python software libraries."

    add_section_heading("Theory")
    for paragraph in theory_text.split("\n\n"):
        if paragraph.strip():
            add_body_text(paragraph.strip())

    add_section_heading("Software Used")
    add_body_text("Python 3.11, Docker Environment, CrewAI Agent Framework.")

    add_section_heading("Code")
    # Source Code Block -> Times New Roman, 14pt, Regular, Left (Preserves Indentation)
    for line in code_content.splitlines():
        p_line = doc.add_paragraph()
        p_line.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p_line.add_run(line if line else " ")
        format_run(r, font_name="Times New Roman", size_pt=14, bold=False)

    add_section_heading("Output")
    add_body_text("Terminal Output Screenshot:")
    terminal_img = output_dir / "terminal_output.png"
    if terminal_img.is_file():
        doc.add_picture(str(terminal_img), width=Inches(6.0))

    plot_files = [p for p in output_dir.glob("*.png") if p.name != "terminal_output.png"]
    if plot_files:
        add_body_text("Generated Visualizations & Plots:")
        for plot_img in plot_files:
            doc.add_picture(str(plot_img), width=Inches(6.0))

    target_doc = output_dir / "solution.docx"
    try:
        doc.save(target_doc)
    except PermissionError:
        target_doc = output_dir / "solution_latest.docx"
        doc.save(target_doc)


sandbox_tools = [
    list_sandbox_files,
    read_sandbox_file,
    write_sandbox_file,
    run_sandbox_python,
]


def _never_cache(*_args, **_kwargs) -> bool:
    return False


# Sandbox state changes between calls (files appear/change/run), so caching tool
# results would feed agents stale data. Opt out of CrewAI's default tool caching.
for _t in sandbox_tools:
    _t.cache_function = _never_cache
